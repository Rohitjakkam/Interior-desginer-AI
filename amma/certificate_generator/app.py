"""
Certificate Bulk Generator
A Streamlit application for generating bulk certificates from Word templates
with automatic serial number incrementing.
"""

import streamlit as st
import re
import io
import os
import zipfile
import tempfile
from docx import Document
from docx.shared import Pt
from typing import List, Dict, Tuple, Optional


def is_doc_support_available() -> bool:
    """Check if .doc file support is available (Windows with pywin32 only)."""
    try:
        import win32com.client
        import pythoncom
        return True
    except ImportError:
        return False


DOC_SUPPORT_AVAILABLE = is_doc_support_available()


def convert_doc_to_docx(doc_bytes: bytes) -> bytes:
    """Convert .doc file to .docx format using Microsoft Word (Windows only).

    Returns the .docx file as bytes.
    """
    try:
        import win32com.client
        import pythoncom

        # Initialize COM for this thread
        pythoncom.CoInitialize()

        try:
            # Create temporary files
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_doc:
                tmp_doc.write(doc_bytes)
                tmp_doc_path = tmp_doc.name

            tmp_docx_path = tmp_doc_path.replace('.doc', '.docx')

            # Use Word to convert
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            try:
                doc = word.Documents.Open(tmp_doc_path)
                # Save as .docx (format 16 = wdFormatDocumentDefault for .docx)
                doc.SaveAs2(tmp_docx_path, FileFormat=16)
                doc.Close()
            finally:
                word.Quit()

            # Read the converted file
            with open(tmp_docx_path, 'rb') as f:
                docx_bytes = f.read()

            # Clean up temp files
            try:
                os.unlink(tmp_doc_path)
                os.unlink(tmp_docx_path)
            except:
                pass

            return docx_bytes

        finally:
            pythoncom.CoUninitialize()

    except ImportError:
        raise Exception(
            "To support .doc files, please install pywin32:\n"
            "pip install pywin32\n\n"
            "Note: This requires Microsoft Word to be installed on your computer."
        )
    except Exception as e:
        raise Exception(f"Failed to convert .doc file: {str(e)}\n\nMake sure Microsoft Word is installed.")


def extract_text_with_positions(doc: Document) -> List[Dict]:
    """Extract all text from document with paragraph and run positions."""
    text_items = []

    for para_idx, paragraph in enumerate(doc.paragraphs):
        full_para_text = paragraph.text
        if full_para_text.strip():
            text_items.append({
                'type': 'paragraph',
                'para_idx': para_idx,
                'text': full_para_text,
                'location': f"Paragraph {para_idx + 1}"
            })

    # Also check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    text_items.append({
                        'type': 'table_cell',
                        'table_idx': table_idx,
                        'row_idx': row_idx,
                        'cell_idx': cell_idx,
                        'text': cell_text,
                        'location': f"Table {table_idx + 1}, Row {row_idx + 1}, Cell {cell_idx + 1}"
                    })

    return text_items


def find_serial_patterns(text: str) -> List[Dict]:
    """Find potential serial number patterns in text."""
    patterns = []

    # Pattern 1: Numbers at end after slash (e.g., /014501)
    matches = re.finditer(r'/(\d{3,})\b', text)
    for match in matches:
        patterns.append({
            'full_match': match.group(0),
            'number': match.group(1),
            'start': match.start(),
            'end': match.end(),
            'pattern_type': 'slash_number'
        })

    # Pattern 2: Standalone numbers (e.g., 014501)
    matches = re.finditer(r'\b(\d{4,})\b', text)
    for match in matches:
        # Avoid duplicates from slash pattern
        already_found = any(p['number'] == match.group(1) for p in patterns)
        if not already_found:
            patterns.append({
                'full_match': match.group(0),
                'number': match.group(1),
                'start': match.start(),
                'end': match.end(),
                'pattern_type': 'standalone_number'
            })

    return patterns


def detect_serial_fields(doc: Document) -> List[Dict]:
    """Automatically detect potential serial number fields in the document."""
    detected_fields = []
    text_items = extract_text_with_positions(doc)

    keywords = ['certificate', 'serial', 'number', 'identification', 'id', 'no.', 'no:', 'ref']

    for item in text_items:
        text_lower = item['text'].lower()

        # Check if line contains serial-related keywords
        has_keyword = any(kw in text_lower for kw in keywords)

        # Find number patterns
        patterns = find_serial_patterns(item['text'])

        if patterns and has_keyword:
            for pattern in patterns:
                detected_fields.append({
                    **item,
                    'pattern': pattern,
                    'suggested': True
                })
        elif patterns:
            for pattern in patterns:
                detected_fields.append({
                    **item,
                    'pattern': pattern,
                    'suggested': False
                })

    return detected_fields


def increment_serial(serial: str, increment: int, preserve_length: bool = True) -> str:
    """Increment a serial number string while preserving format."""
    # Extract numeric part
    num = int(serial)
    new_num = num + increment

    if preserve_length:
        return str(new_num).zfill(len(serial))
    return str(new_num)


def replace_last_occurrence(text: str, old: str, new: str) -> str:
    """Replace only the last occurrence of 'old' with 'new' in 'text'."""
    # Find the last occurrence
    last_pos = text.rfind(old)
    if last_pos == -1:
        return text
    return text[:last_pos] + new + text[last_pos + len(old):]


def replace_in_paragraph(paragraph, old_text: str, new_text: str):
    """Replace text in paragraph while preserving formatting.

    This function carefully handles text replacement to maintain
    the original font size, style, and other formatting.
    """
    full_text = paragraph.text
    if old_text not in full_text:
        return

    runs = paragraph.runs
    if not runs:
        return

    # Simple case: text is entirely within a single run
    for run in runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return

    # Complex case: text spans multiple runs
    # Find where the old_text starts and ends across runs
    start_pos = full_text.find(old_text)
    end_pos = start_pos + len(old_text)

    # Build a map of run boundaries
    run_boundaries = []  # [(start, end, run), ...]
    current_pos = 0
    for run in runs:
        run_start = current_pos
        run_end = current_pos + len(run.text)
        run_boundaries.append((run_start, run_end, run))
        current_pos = run_end

    # Find which runs are affected
    for i, (run_start, run_end, run) in enumerate(run_boundaries):
        # Check if this run contains the start of old_text
        if run_start <= start_pos < run_end:
            # This run contains the start of the text to replace
            text_before = run.text[:start_pos - run_start]

            # Check if the entire old_text is within this run
            if end_pos <= run_end:
                # Simple case within one run
                text_after = run.text[end_pos - run_start:]
                run.text = text_before + new_text + text_after
                return
            else:
                # Text spans multiple runs
                # Put the replacement in this run with text before it
                run.text = text_before + new_text

                # Clear text from subsequent runs that are part of old_text
                remaining_to_clear = end_pos - run_end
                for j in range(i + 1, len(run_boundaries)):
                    _, _, next_run = run_boundaries[j]
                    next_run_len = len(next_run.text)

                    if remaining_to_clear >= next_run_len:
                        # Clear entire run
                        next_run.text = ""
                        remaining_to_clear -= next_run_len
                    else:
                        # Partial clear - keep the rest
                        next_run.text = next_run.text[remaining_to_clear:]
                        remaining_to_clear = 0
                        break
                return


def replace_in_table_cell(cell, old_text: str, new_text: str):
    """Replace text in a table cell while preserving formatting."""
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, old_text, new_text)


def remove_highlighting(doc: Document):
    """Remove all yellow/any highlighting from the document."""
    from docx.oxml.ns import qn

    # Remove highlighting from paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = None
            # Also remove via XML if needed
            rPr = run._r.get_or_add_rPr()
            highlight = rPr.find(qn('w:highlight'))
            if highlight is not None:
                rPr.remove(highlight)

    # Remove highlighting from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.highlight_color = None
                        rPr = run._r.get_or_add_rPr()
                        highlight = rPr.find(qn('w:highlight'))
                        if highlight is not None:
                            rPr.remove(highlight)


def generate_certificate(template_bytes: bytes, field_mappings: List[Dict], increment: int) -> Document:
    """Generate a single certificate with incremented serial numbers."""
    # Create a new document from template bytes
    new_doc = Document(io.BytesIO(template_bytes))

    # Remove all highlighting from the document
    remove_highlighting(new_doc)

    for mapping in field_mappings:
        old_text = mapping['pattern']['full_match']

        # Check if this is a manual field with multiple numbers
        if mapping['type'] == 'manual' and 'numbers' in mapping['pattern']:
            numbers_list = mapping['pattern']['numbers']
            new_text = old_text

            # Replace each number from right to left (to handle overlapping positions)
            # We process in reverse order of position in the string
            number_positions = []
            for num in numbers_list:
                pos = new_text.rfind(num)
                if pos != -1:
                    number_positions.append((pos, num))

            # Sort by position descending (right to left)
            number_positions.sort(key=lambda x: x[0], reverse=True)

            # Replace each number
            for pos, num in number_positions:
                new_num = increment_serial(num, increment)
                new_text = new_text[:pos] + new_num + new_text[pos + len(num):]

            # Search and replace in entire document
            for paragraph in new_doc.paragraphs:
                if old_text in paragraph.text:
                    replace_in_paragraph(paragraph, old_text, new_text)

            for table in new_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_text in cell.text:
                            replace_in_table_cell(cell, old_text, new_text)

        else:
            # Standard single number replacement
            old_serial = mapping['pattern']['number']
            new_serial = increment_serial(old_serial, increment)
            # Use replace_last_occurrence to handle cases like "1687/2526/1"
            # where we want to replace only the last "1", not the "1" in "1687"
            new_text = replace_last_occurrence(old_text, old_serial, new_serial)

            if mapping['type'] == 'paragraph':
                para_idx = mapping['para_idx']
                if para_idx < len(new_doc.paragraphs):
                    replace_in_paragraph(new_doc.paragraphs[para_idx], old_text, new_text)

            elif mapping['type'] == 'table_cell':
                table_idx = mapping['table_idx']
                row_idx = mapping['row_idx']
                cell_idx = mapping['cell_idx']

                if table_idx < len(new_doc.tables):
                    table = new_doc.tables[table_idx]
                    if row_idx < len(table.rows):
                        row = table.rows[row_idx]
                        if cell_idx < len(row.cells):
                            replace_in_table_cell(row.cells[cell_idx], old_text, new_text)

            elif mapping['type'] == 'manual':
                # For manual fields with single number, search entire document
                for paragraph in new_doc.paragraphs:
                    if old_text in paragraph.text:
                        replace_in_paragraph(paragraph, old_text, new_text)

                for table in new_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if old_text in cell.text:
                                replace_in_table_cell(cell, old_text, new_text)

    return new_doc


def save_doc_to_bytes(doc: Document) -> bytes:
    """Save a document to bytes."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def remove_blank_pages(doc: Document):
    """Remove blank pages from the document.

    This function removes:
    - Empty paragraphs that only contain page breaks
    - Paragraphs with only whitespace
    - Consecutive empty paragraphs that create blank pages
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    elements_to_remove = []

    # Get all paragraph elements
    paragraphs = body.findall(qn('w:p'))

    for i, para in enumerate(paragraphs):
        # Check if paragraph is empty or only has formatting
        text_content = ''.join(para.itertext()).strip()

        # Check if it has a page break but no real content
        has_page_break = False
        has_content = bool(text_content)

        # Check for page break elements
        for br in para.iter(qn('w:br')):
            br_type = br.get(qn('w:type'))
            if br_type == 'page':
                has_page_break = True

        # Check for pageBreakBefore in paragraph properties
        pPr = para.find(qn('w:pPr'))
        if pPr is not None:
            pb_before = pPr.find(qn('w:pageBreakBefore'))
            if pb_before is not None:
                has_page_break = True

        # If paragraph has page break but no content, and is followed by another
        # element with page break, mark for removal
        if has_page_break and not has_content:
            # Check if next paragraph also starts with page break
            if i + 1 < len(paragraphs):
                next_para = paragraphs[i + 1]
                next_pPr = next_para.find(qn('w:pPr'))
                if next_pPr is not None:
                    next_pb = next_pPr.find(qn('w:pageBreakBefore'))
                    if next_pb is not None:
                        # This empty page break paragraph creates a blank page
                        elements_to_remove.append(para)
                        continue

        # Remove completely empty paragraphs (no text, no images, no tables)
        if not has_content and not has_page_break:
            # Check if it has any runs with content
            runs = para.findall(qn('w:r'))
            has_any_content = False
            for run in runs:
                # Check for drawings/images
                if run.find(qn('w:drawing')) is not None:
                    has_any_content = True
                    break
                # Check for any text
                for t in run.iter(qn('w:t')):
                    if t.text and t.text.strip():
                        has_any_content = True
                        break

            # If truly empty, check if it's creating a blank page
            if not has_any_content:
                # Check previous element for page break
                prev_idx = list(body).index(para) - 1
                if prev_idx >= 0:
                    prev_elem = list(body)[prev_idx]
                    if prev_elem.tag.endswith('}p'):
                        prev_text = ''.join(prev_elem.itertext()).strip()
                        if not prev_text:
                            # Two consecutive empty paragraphs - remove one
                            elements_to_remove.append(para)

    # Remove marked elements
    for elem in elements_to_remove:
        try:
            body.remove(elem)
        except ValueError:
            pass  # Element already removed

    # Second pass: remove standalone empty page break paragraphs at problematic positions
    paragraphs = body.findall(qn('w:p'))
    for para in paragraphs:
        text_content = ''.join(para.itertext()).strip()
        if not text_content:
            # Check if this paragraph only has a page break and nothing else useful
            runs = para.findall(qn('w:r'))
            only_has_break = False

            for run in runs:
                br_elements = run.findall(qn('w:br'))
                for br in br_elements:
                    if br.get(qn('w:type')) == 'page':
                        only_has_break = True

                # If run has other content, don't remove
                if run.find(qn('w:drawing')) is not None:
                    only_has_break = False
                    break

            if only_has_break and len(runs) <= 1:
                # This is likely causing a blank page
                try:
                    body.remove(para)
                except ValueError:
                    pass


def create_combined_document(template_bytes: bytes, field_mappings: List[Dict], count: int) -> bytes:
    """Create a single document with all certificates separated by page breaks.

    This function properly preserves the template formatting by copying XML elements
    from each generated certificate into a combined document.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    # Generate the first certificate as the base document
    combined_doc = generate_certificate(template_bytes, field_mappings, 0)

    # For remaining certificates, append with page breaks
    for i in range(1, count):
        # Generate each certificate from template
        cert_doc = generate_certificate(template_bytes, field_mappings, i)

        # Get all body elements from cert_doc (excluding sectPr)
        elements_to_copy = []
        for element in cert_doc.element.body:
            if not element.tag.endswith('sectPr'):
                elements_to_copy.append(deepcopy(element))

        if elements_to_copy:
            # Add page break to the first element of this certificate
            first_element = elements_to_copy[0]

            # Check if it's a paragraph (w:p) - add page break before
            if first_element.tag.endswith('}p'):
                # Find or create pPr (paragraph properties)
                pPr = first_element.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    first_element.insert(0, pPr)

                # Add page break before this paragraph
                pageBreakBefore = OxmlElement('w:pageBreakBefore')
                pageBreakBefore.set(qn('w:val'), 'true')
                pPr.append(pageBreakBefore)

            elif first_element.tag.endswith('}tbl'):
                # For tables, insert a page break paragraph before the table
                page_break_para = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                pageBreakBefore = OxmlElement('w:pageBreakBefore')
                pageBreakBefore.set(qn('w:val'), 'true')
                pPr.append(pageBreakBefore)
                page_break_para.append(pPr)
                combined_doc.element.body.append(page_break_para)

            # Append all elements
            for element in elements_to_copy:
                combined_doc.element.body.append(element)

    # Remove any blank pages that may have been created
    remove_blank_pages(combined_doc)

    return save_doc_to_bytes(combined_doc)


def main():
    st.set_page_config(
        page_title="Certificate Bulk Generator",
        page_icon="📜",
        layout="wide"
    )

    st.title("📜 Certificate Bulk Generator")
    st.markdown("Generate multiple certificates from a Word template with automatic serial number incrementing.")

    # Initialize session state
    if 'template_doc' not in st.session_state:
        st.session_state.template_doc = None
    if 'template_bytes' not in st.session_state:
        st.session_state.template_bytes = None
    if 'selected_fields' not in st.session_state:
        st.session_state.selected_fields = []

    # Sidebar for instructions
    with st.sidebar:
        st.header("Instructions")
        doc_formats = ".doc or .docx" if DOC_SUPPORT_AVAILABLE else ".docx"
        st.markdown(f"""
        1. **Upload Template**: Upload your Word document ({doc_formats})
        2. **Add Fields**: Enter the text patterns with serial numbers to increment
        3. **Configure**: Set the number of certificates to generate
        4. **Generate**: Download individual files or a combined document
        """)

        st.header("Example")
        st.markdown("""
        **Text to search:** `1687/2526/1`

        **Numbers to increment:** `1` (only last number)

        Or: `2526,1` (both numbers)
        """)

        st.header("Supported Formats")
        if DOC_SUPPORT_AVAILABLE:
            st.markdown("""
            - `.docx` (Word 2007+)
            - `.doc` (Legacy Word)
            """)
        else:
            st.markdown("""
            - `.docx` (Word 2007+)

            *.doc support requires Windows + MS Word*
            """)

    # Step 1: Upload Template
    st.header("Step 1: Upload Template")

    # Conditionally support .doc files based on platform
    if DOC_SUPPORT_AVAILABLE:
        uploaded_file = st.file_uploader(
            "Upload your Word document template (.doc or .docx)",
            type=['docx', 'doc'],
            help="Upload a Word document that will serve as your certificate template. Both .doc and .docx formats are supported."
        )
    else:
        uploaded_file = st.file_uploader(
            "Upload your Word document template (.docx)",
            type=['docx'],
            help="Upload a Word document (.docx) that will serve as your certificate template."
        )

    if uploaded_file is not None:
        # Load document
        try:
            file_bytes = uploaded_file.getvalue()
            filename = uploaded_file.name.lower()

            # Check if it's a .doc file and convert it (only on Windows with pywin32)
            if DOC_SUPPORT_AVAILABLE and filename.endswith('.doc') and not filename.endswith('.docx'):
                with st.spinner("Converting .doc to .docx format..."):
                    st.session_state.template_bytes = convert_doc_to_docx(file_bytes)
                st.info("📄 Converted .doc file to .docx format")
            else:
                st.session_state.template_bytes = file_bytes

            st.session_state.template_doc = Document(io.BytesIO(st.session_state.template_bytes))
            st.success(f"✅ Template loaded: {uploaded_file.name}")

        except Exception as e:
            st.error(f"Error loading document: {str(e)}")
            st.session_state.template_doc = None

    if st.session_state.template_doc is not None:
        # Step 2: Field Mapping
        st.header("Step 2: Add Custom Fields to Increment")

        # Initialize selected_fields
        st.session_state.selected_fields = []

        st.markdown("Add the text patterns containing serial numbers that should be incremented. You can increment **multiple numbers** in the same text.")

        with st.expander("Add custom field", expanded=True):
            manual_search_text = st.text_input(
                "Text to search for (e.g., '1687/2526/1')",
                key="manual_search",
                help="Enter the exact text pattern including the numbers that should increment"
            )
            manual_numbers = st.text_input(
                "Numbers to increment (comma-separated, e.g., '2526,1')",
                key="manual_numbers",
                help="Enter the numeric parts that should increment, separated by commas. Each will increment by 1 for each certificate."
            )

            if st.button("Add Field", key="add_manual"):
                if manual_search_text and manual_numbers:
                    # Parse comma-separated numbers
                    numbers_list = [n.strip() for n in manual_numbers.split(',') if n.strip()]

                    # Validate all numbers are in the search text
                    invalid_numbers = [n for n in numbers_list if n not in manual_search_text]

                    if invalid_numbers:
                        st.error(f"These numbers are not found in the search text: {', '.join(invalid_numbers)}")
                    elif not numbers_list:
                        st.error("Please enter at least one number to increment.")
                    else:
                        manual_field = {
                            'type': 'manual',
                            'text': f"Manual: {manual_search_text}",
                            'location': "User-defined",
                            'pattern': {
                                'full_match': manual_search_text,
                                'number': numbers_list[0],  # Primary number for compatibility
                                'numbers': numbers_list,  # All numbers to increment
                                'pattern_type': 'manual'
                            },
                            'suggested': False
                        }
                        if 'manual_fields' not in st.session_state:
                            st.session_state.manual_fields = []
                        st.session_state.manual_fields.append(manual_field)
                        st.success(f"Added field: {manual_search_text} (incrementing: {', '.join(numbers_list)})")
                        st.rerun()
                else:
                    st.error("Please fill in both fields.")

        # Display and select custom fields
        if 'manual_fields' in st.session_state and st.session_state.manual_fields:
            st.subheader("Added Fields")
            for idx, field in enumerate(st.session_state.manual_fields):
                col1, col2, col3, col4 = st.columns([1, 3, 2, 1])
                with col1:
                    if st.checkbox("Select", key=f"manual_{idx}", value=True):
                        st.session_state.selected_fields.append(field)
                with col2:
                    st.text(field['pattern']['full_match'])
                with col3:
                    numbers = field['pattern'].get('numbers', [field['pattern']['number']])
                    st.caption(f"Incrementing: **{', '.join(numbers)}**")
                with col4:
                    if st.button("🗑️", key=f"delete_manual_{idx}"):
                        st.session_state.manual_fields.pop(idx)
                        st.rerun()

        # Step 3: Generation Settings
        st.header("Step 3: Configure Generation")

        col1, col2 = st.columns(2)

        with col1:
            num_certificates = st.number_input(
                "Number of certificates to generate",
                min_value=1,
                max_value=20000,
                value=10,
                help="How many certificates do you want to generate?"
            )

        with col2:
            output_format = st.radio(
                "Output format",
                options=["Individual files (ZIP)", "Combined document"],
                help="Choose how you want to receive the generated certificates"
            )

        # Preview
        if st.session_state.selected_fields:
            st.subheader("Preview")
            st.markdown("**Selected fields will increment as follows:**")

            def generate_preview_text(field, increment_val):
                """Generate preview text for a field with given increment."""
                full_match = field['pattern']['full_match']

                # Check if multiple numbers
                if 'numbers' in field['pattern']:
                    numbers_list = field['pattern']['numbers']
                    result = full_match

                    # Get positions and replace from right to left
                    number_positions = []
                    for num in numbers_list:
                        pos = result.rfind(num)
                        if pos != -1:
                            number_positions.append((pos, num))

                    number_positions.sort(key=lambda x: x[0], reverse=True)

                    for pos, num in number_positions:
                        new_num = increment_serial(num, increment_val)
                        result = result[:pos] + new_num + result[pos + len(num):]

                    return result
                else:
                    serial = field['pattern']['number']
                    return replace_last_occurrence(full_match, serial, increment_serial(serial, increment_val))

            preview_data = []
            for field in st.session_state.selected_fields:
                full_match = field['pattern']['full_match']
                preview_data.append({
                    'Pattern': full_match,
                    'Cert #1': generate_preview_text(field, 0),
                    'Cert #2': generate_preview_text(field, 1),
                    'Cert #3': generate_preview_text(field, 2),
                    f'Cert #{num_certificates}': generate_preview_text(field, num_certificates - 1)
                })

            st.table(preview_data)

        # Step 4: Generate
        st.header("Step 4: Generate Certificates")

        if not st.session_state.selected_fields:
            st.warning("Please select at least one field to increment.")
        else:
            if st.button("🚀 Generate Certificates", type="primary", use_container_width=True):
                with st.spinner(f"Generating {num_certificates} certificates..."):
                    try:
                        progress_bar = st.progress(0)

                        if output_format == "Individual files (ZIP)":
                            # Create ZIP file with individual documents
                            zip_buffer = io.BytesIO()

                            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                                for i in range(num_certificates):
                                    # Generate certificate from template bytes
                                    cert_doc = generate_certificate(
                                        st.session_state.template_bytes,
                                        st.session_state.selected_fields,
                                        i
                                    )

                                    # Get first serial for filename
                                    first_serial = st.session_state.selected_fields[0]['pattern']['number']
                                    cert_serial = increment_serial(first_serial, i)

                                    # Save to bytes
                                    doc_bytes = save_doc_to_bytes(cert_doc)
                                    zip_file.writestr(f"Certificate_{cert_serial}.docx", doc_bytes)

                                    progress_bar.progress((i + 1) / num_certificates)

                            zip_buffer.seek(0)

                            st.success(f"✅ Generated {num_certificates} certificates!")

                            st.download_button(
                                label="📥 Download ZIP File",
                                data=zip_buffer.getvalue(),
                                file_name="certificates.zip",
                                mime="application/zip",
                                use_container_width=True
                            )

                        else:
                            # Create combined document
                            combined_bytes = create_combined_document(
                                st.session_state.template_bytes,
                                st.session_state.selected_fields,
                                num_certificates
                            )

                            progress_bar.progress(1.0)

                            st.success(f"✅ Generated combined document with {num_certificates} certificates!")

                            st.download_button(
                                label="📥 Download Combined Document",
                                data=combined_bytes,
                                file_name="certificates_combined.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )

                    except Exception as e:
                        st.error(f"Error generating certificates: {str(e)}")
                        import traceback
                        st.code(traceback.format_exc())

    # Footer
    st.markdown("---")
    st.markdown(
        "<div style='text-align: center; color: gray;'>"
        "Certificate Bulk Generator | Made with Streamlit"
        "</div>",
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()
